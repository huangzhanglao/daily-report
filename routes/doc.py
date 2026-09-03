# -*- coding: utf-8 -*-
"""AI 文档工作台 —— APIRouter。

能力：接入大模型帮你写材料（复用「设置 → 大模型配置」里每个用户自己接入的 OpenAI 兼容服务）。
写作流程内置「反思」机制，共三次模型调用：
  1) 生成初稿   —— 注入文体规范 + 素材 + 需求
  2) 反思审稿   —— 让模型扮演资深审稿人，指出初稿的问题与改进点
  3) 反思改写   —— 依据审稿意见把初稿改写为高质量终稿

输入素材支持两种：直接粘贴一段文字；或上传文件（.txt/.md 原生，.docx/.pdf 依赖可选解析库，缺失时友好提示）。

鉴权：require_uid，所有数据按 user_id 隔离。
"""
import json
import time
import httpx
from fastapi import APIRouter, Depends, HTTPException, Request, Header, Response, UploadFile, File
from fastapi.responses import JSONResponse

from core import (
    get_conn, write_lock, uid_now, row_to_user, blog_public,
    get_uid, require_uid, require_admin,
    sign_token, verify_password, hash_password,
    set_auth_cookie, clear_auth_cookie, AUTH_COOKIE,
    rate_limit, client_ip, public_register_enabled,
    grant_all_apps, grant_app_to_all_users,
    DOC_TYPES,
)

router = APIRouter()

# ---------------- 行映射 ----------------
def _doc_row(r) -> dict:
    return {
        "id": r["id"], "doc_type": r["doc_type"], "title": r["title"],
        "requirement": r["requirement"], "material": r["material"],
        "draft": r["draft"], "critique": r["critique"], "final": r["final"],
        "status": r["status"],
        "createdAt": r["created_at"], "updatedAt": r["updated_at"],
    }

def _doc_meta(r) -> dict:  # 列表用的精简行（不含长正文，省流量）
    return {
        "id": r["id"], "doc_type": r["doc_type"], "title": r["title"],
        "status": r["status"],
        "createdAt": r["created_at"], "updatedAt": r["updated_at"],
    }

# ---------------- 文体预设辅助 ----------------
def _doc_system(doc_type: str) -> str:
    for t in DOC_TYPES:
        if t["key"] == doc_type:
            return t["system"]
    return DOC_TYPES[0]["system"]

@router.get("/api/doc/types")
def list_doc_types(uid: str = Depends(require_uid)):
    """返回文体预设（key/name/icon/short），供前端下拉。"""
    return {"types": DOC_TYPES}

# ---------------- 大模型调用（复用 llm_providers 配置，逻辑同 llm.py） ----------------
def _normalize_base_url(u: str) -> str:
    u = (u or "").strip()
    if not u:
        return ""
    if not u.startswith("http://") and not u.startswith("https://"):
        u = "https://" + u
    return u.rstrip("/")

def _pick_provider(conn, uid: str, pid: str = ""):
    """选 provider：优先指定 id → 否则默认 → 否则最近更新的一个。返回行或 None。"""
    if pid:
        r = conn.execute("SELECT * FROM llm_providers WHERE id=? AND user_id=?", (pid, uid)).fetchone()
        if r:
            return r
    r = conn.execute("SELECT * FROM llm_providers WHERE user_id=? AND is_default=1 LIMIT 1", (uid,)).fetchone()
    if r:
        return r
    return conn.execute("SELECT * FROM llm_providers WHERE user_id=? ORDER BY updated_at DESC LIMIT 1", (uid,)).fetchone()

async def _chat(uid: str, messages, temperature: float = 0.7, max_tokens: int = 4000, pid: str = ""):
    """调用大模型 chat/completions，返回回复文本。无可用配置时抛 400。"""
    conn = get_conn()
    try:
        r = _pick_provider(conn, uid, pid)
    finally:
        conn.close()
    if not r:
        raise HTTPException(400, "尚未配置任何大模型，请先在「设置 → 大模型配置」中添加")
    base_url = _normalize_base_url(r["base_url"])
    payload = {
        "model": r["model"], "messages": messages,
        "temperature": temperature, "max_tokens": max_tokens,
    }
    url = base_url + "/chat/completions"
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            resp = await client.post(url, headers={"Authorization": "Bearer " + r["api_key"], "Content-Type": "application/json"}, json=payload)
            data = resp.json()
    except Exception as e:
        raise HTTPException(502, "调用大模型失败：" + str(e)[:200])
    if resp.status_code != 200:
        raise HTTPException(resp.status_code, "大模型返回错误：" + str(data)[:300])
    try:
        return data["choices"][0]["message"]["content"]
    except Exception:
        raise HTTPException(502, "大模型返回格式异常")

# ---------------- 文件素材抽取 ----------------
ALLOWED_EXTS = {".txt", ".md", ".docx", ".pdf"}
MAX_UPLOAD_BYTES = 3 * 1024 * 1024  # 3MB

def _extract_text(ext: str, raw: bytes) -> str:
    ext = ext.lower()
    if ext in (".txt", ".md"):
        # 多种编码尝试，保证中文不乱码
        for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
            try:
                return raw.decode(enc)
            except Exception:
                continue
        return raw.decode("utf-8", errors="ignore")
    if ext == ".docx":
        try:
            import docx  # 延迟导入：python-docx
        except Exception:
            raise HTTPException(400, "解析 .docx 需要服务端安装 python-docx，当前缺失；请改用 .txt/.md 上传")
        try:
            d = docx.Document(__import__("io").BytesIO(raw))
            parts = [p.text for p in d.paragraphs]
            for tb in d.tables:
                for row in tb.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(x for x in parts if x and x.strip())
        except Exception as e:
            raise HTTPException(400, "解析 .docx 失败：" + str(e)[:200])
    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # 延迟导入：pypdf
        except Exception:
            raise HTTPException(400, "解析 .pdf 需要服务端安装 pypdf，当前缺失；请改用 .txt/.md 上传")
        try:
            reader = PdfReader(__import__("io").BytesIO(raw))
            parts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            if not parts:
                raise HTTPException(400, "该 PDF 无可提取文本（可能是扫描件图片，暂不支持 OCR）")
            return "\n".join(parts)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, "解析 .pdf 失败：" + str(e)[:200])
    raise HTTPException(400, "暂不支持该文件类型，仅支持 .txt / .md / .docx / .pdf")

@router.post("/api/doc/upload")
async def doc_upload(file: UploadFile = File(...), uid: str = Depends(require_uid)):
    """上传素材文件并抽取为纯文本返回（不落盘）。"""
    name = (file.filename or "").strip()
    ext = "." + (name.rsplit(".", 1)[-1].lower() if "." in name else "")
    if ext not in ALLOWED_EXTS:
        raise HTTPException(400, "仅支持 .txt / .md / .docx / .pdf 文件")
    raw = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(400, "文件超过 3MB 限制")
    try:
        text = _extract_text(ext, raw)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, "读取文件失败：" + str(e)[:200])
    # 裁剪到合理长度（避免把素材一股脑塞给模型超出上下文）
    text = text.strip()
    if len(text) > 60000:
        text = text[:60000] + "\n…（素材过长已截断）"
    return {"ok": True, "name": name, "ext": ext, "text": text, "chars": len(text)}

# ---------------- 反思式写作管线 ----------------
def _build_user_prompt(doc_type: str, title: str, requirement: str, material: str) -> str:
    p = []
    if doc_type and doc_type != "general":
        p.append("请以文体要求写作。")
    if title:
        p.append("文档标题（或主题）：" + title)
    if requirement:
        p.append("写作需求：" + requirement)
    if material and material.strip():
        p.append("\n【可参考素材如下】\n" + material.strip())
    if not requirement and not material:
        p.append("请直接输出一篇完整的文档。")
    p.append("\n请直接输出文档正文，不要输出额外说明。")
    return "\n".join(p)

async def _gen_draft(uid: str, doc_type: str, title: str, requirement: str, material: str, pid: str):
    system = _doc_system(doc_type)
    msgs = [
        {"role": "system", "content": system},
        {"role": "user", "content": _build_user_prompt(doc_type, title, requirement, material)},
    ]
    return (await _chat(uid, msgs, temperature=0.7, max_tokens=4000, pid=pid)).strip()

async def _gen_critique(uid: str, doc_type: str, title: str, requirement: str, material: str, draft: str, pid: str):
    """让模型以资深审稿人身份评估初稿，输出结构化意见，供改写阶段使用。"""
    sys_prompt = (
        "你是一位极其严格、专业的中文审稿人，精通各类公文与书面材料的写作规范。"
        "请对用户给出的『初稿』进行严格审视，从以下维度逐条指出真实存在的问题："
        "①结构与逻辑 ②语言与表达 ③内容完整性（是否覆盖需求/素材关键点）④文体规范性 ⑤细节（错别字/标点/数据前后矛盾）。\n"
        "要求：只做建设性批评，具体到『哪一段/哪一句』有何问题、应如何改进；不要空泛表扬。\n"
        "请用如下结构输出，便于机器与人类阅读：\n"
        "【总体评价】一两句话\n"
        "【具体问题】\n- 问题1…（定位+原因+改法）\n- 问题2…\n"
        "【改进要求】明确列出改写时必须落实的要点（供下一步改写严格照做）"
    )
    ctx = ""
    if title:
        ctx += "文档主题：" + title + "\n"
    if requirement:
        ctx += "写作需求：" + requirement + "\n"
    if material and material.strip():
        ctx += "参考素材：" + material.strip()[:3000] + "\n"
    msgs = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": (ctx + "\n【初稿】\n" + draft)},
    ]
    return (await _chat(uid, msgs, temperature=0.2, max_tokens=2000, pid=pid)).strip()

async def _gen_final(uid: str, doc_type: str, title: str, requirement: str, material: str, draft: str, critique: str, pid: str):
    """依据审稿意见把初稿改写为高质量终稿。"""
    sys_prompt = (
        "你是一位顶尖中文写作专家。下面给你一份『初稿』和资深审稿人给出的『审稿意见』。"
        "请严格依据审稿意见，对初稿进行改写优化，产出一篇质量明显更高的终稿："
        "1. 必须落实审稿意见中的每一条『改进要求』；"
        "2. 保持文体规范和用户需求的完整覆盖，素材中的关键信息不得遗漏或篡改；"
        "3. 语言精炼、逻辑严谨、层次清晰；"
        "4. 直接输出改写后的完整终稿正文，不要复述意见，不要加任何前言或说明。"
    )
    msgs = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": (
            "文体要求：" + _doc_system(doc_type)[:1500] + "\n\n"
            "写作需求：" + (requirement or "") + "\n\n"
            "【初稿】\n" + draft + "\n\n"
            "【审稿意见】\n" + critique
        )},
    ]
    return (await _chat(uid, msgs, temperature=0.5, max_tokens=4000, pid=pid)).strip()

@router.post("/api/doc/generate")
async def doc_generate(body: dict, uid: str = Depends(require_uid)):
    """反思式写作：初稿 → 审稿 → 改写。body: {doc_type,title,requirement,material,provider_id,doc_id?}。
    若提供 doc_id 则更新已有文档；否则新建。"""
    doc_type = (body.get("doc_type") or "general").strip()
    title = (body.get("title") or "").strip()
    requirement = (body.get("requirement") or "").strip()
    material = (body.get("material") or "").strip()
    pid = (body.get("provider_id") or "").strip()
    if not title and not requirement:
        raise HTTPException(400, "请至少填写文档标题或写作需求")
    # 三步顺序调用（任一步失败即返回当前进度）
    draft = await _gen_draft(uid, doc_type, title, requirement, material, pid)
    critique = await _gen_critique(uid, doc_type, title, requirement, material, draft, pid)
    final = await _gen_final(uid, doc_type, title, requirement, material, draft, critique, pid)

    doc_id = (body.get("doc_id") or "").strip()
    now = int(time.time() * 1000)
    with write_lock:
        conn = get_conn()
        try:
            if doc_id:
                cur = conn.execute("SELECT id FROM doc_docs WHERE id=? AND user_id=?", (doc_id, uid)).fetchone()
                if not cur:
                    raise HTTPException(404, "文档不存在")
                conn.execute(
                    "UPDATE doc_docs SET doc_type=?,title=?,requirement=?,material=?,draft=?,critique=?,final=?,status='done',updated_at=? WHERE id=? AND user_id=?",
                    (doc_type, title, requirement, material, draft, critique, final, now, doc_id, uid),
                )
                out_id = doc_id
            else:
                doc_id = uid_now()
                conn.execute(
                    "INSERT INTO doc_docs (id,user_id,doc_type,title,requirement,material,draft,critique,final,status,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                    (doc_id, uid, doc_type, title, requirement, material, draft, critique, final, "done", now, now),
                )
                out_id = doc_id
            conn.commit()
        finally:
            conn.close()
    return {"ok": True, "id": out_id, "draft": draft, "critique": critique, "final": final, "title": title or "未命名文档"}

@router.post("/api/doc/docs/{did}/regenerate")
async def doc_regenerate(did: str, body: dict, uid: str = Depends(require_uid)):
    """对已有文档再跑一轮『反思改写』（保留原初稿，仅重新审稿并更新终稿）。
    适用于想再打磨一遍的情况。"""
    with write_lock:
        conn = get_conn()
        try:
            r = conn.execute("SELECT * FROM doc_docs WHERE id=? AND user_id=?", (did, uid)).fetchone()
        finally:
            conn.close()
    if not r:
        raise HTTPException(404, "文档不存在")
    pid = (body.get("provider_id") or "").strip()
    doc_type = r["doc_type"]
    critique = await _gen_critique(uid, doc_type, r["title"], r["requirement"], r["material"], r["final"], pid)
    final = await _gen_final(uid, doc_type, r["title"], r["requirement"], r["material"], r["final"], critique, pid)
    with write_lock:
        conn = get_conn()
        try:
            conn.execute(
                "UPDATE doc_docs SET critique=?,final=?,status='done',updated_at=? WHERE id=? AND user_id=?",
                (critique, final, int(time.time() * 1000), did, uid),
            )
            conn.commit()
        finally:
            conn.close()
    return {"ok": True, "id": did, "critique": critique, "final": final}

# ---------------- 文档列表 / 详情 / 删除 ----------------
@router.get("/api/doc/docs")
def doc_list(uid: str = Depends(require_uid), page: int = 1, page_size: int = 50):
    conn = get_conn()
    try:
        total = conn.execute("SELECT COUNT(*) AS c FROM doc_docs WHERE user_id=?", (uid,)).fetchone()["c"]
        page = max(1, page); page_size = max(1, min(100, page_size))
        off = (page - 1) * page_size
        rows = conn.execute(
            "SELECT * FROM doc_docs WHERE user_id=? ORDER BY created_at DESC LIMIT ? OFFSET ?",
            (uid, page_size, off),
        ).fetchall()
    finally:
        conn.close()
    return {"total": total, "items": [_doc_meta(r) for r in rows]}

@router.get("/api/doc/docs/{did}")
def doc_detail(did: str, uid: str = Depends(require_uid)):
    conn = get_conn()
    try:
        r = conn.execute("SELECT * FROM doc_docs WHERE id=? AND user_id=?", (did, uid)).fetchone()
    finally:
        conn.close()
    if not r:
        raise HTTPException(404, "文档不存在")
    return _doc_row(r)

@router.delete("/api/doc/docs/{did}")
def doc_delete(did: str, uid: str = Depends(require_uid)):
    with write_lock:
        conn = get_conn()
        try:
            conn.execute("DELETE FROM doc_docs WHERE id=? AND user_id=?", (did, uid))
            conn.commit()
        finally:
            conn.close()
    return {"ok": True}
